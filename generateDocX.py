from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import json

# Load JSON data from file
with open("output.json", "r", encoding="utf-8") as json_file:
    json_data = json.load(json_file)

# Create a new Word document
doc = Document()



# Iterate through the JSON data and add rows to the table
for item in json_data:

    # Add a table with 2 columns
    table = doc.add_table(rows=0, cols=2)

    # Set column widths
    table.columns[0].width = Pt(15)
    table.columns[1].width = Pt(700)
    table.style = 'Table Grid'
    
    for key, value in item.items():
        if key in ["H", "Đ", "T1", "T2", "T3"]:
            # Add a new row to the table
            row = table.add_row()

            # Format the first column (bold and red color)
            cell_1 = row.cells[0]
            cell_1.text = key
            cell_1.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            cell_1.paragraphs[0].runs[0].font.name = 'HungHau'
            cell_1.paragraphs[0].runs[0].font.size = Pt(12)

            # Add content to the second column
            cell_2 = row.cells[1]
            cell_2.text = value
            cell_2.paragraphs[0].runs[0].font.name = 'HungHau'
            cell_2.paragraphs[0].runs[0].font.size = Pt(12)

            # Check if the key is "H" and modify the second column accordingly
            if key == "H":
                # Create a run for the text before "Câu 1"
                runs = cell_2.paragraphs[0].runs
                runs[0].text = ""  # Get the text before ":"

                # Create a run for "Câu 1" with red color and bold
                run_cau_1 = cell_2.paragraphs[0].add_run(value.split(":")[0] + ": ")
                run_cau_1.font.color.rgb = RGBColor(255, 0, 0)
                run_cau_1.bold = True
                run_cau_1.font.name = 'HungHau'

                # Create a run for the text after "Câu 1"
                run_after_cau_1 = cell_2.paragraphs[0].add_run(value.split(":")[1])
                run_after_cau_1.font.name = 'HungHau'

                
    paragraph = doc.add_paragraph('\n')

# Save the document
doc.save("output.docx")
